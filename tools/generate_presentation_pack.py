from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, ListFlowable, ListItem, PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle

PROJECT_ROOT = Path(__file__).resolve().parents[1]
SRC_ROOT = PROJECT_ROOT / "src"
if str(SRC_ROOT) not in sys.path:
    sys.path.insert(0, str(SRC_ROOT))

from technocops_ddc import APP_NAME, APP_VERSION, APP_VERSION_LABEL, COMPANY_NAME  # noqa: E402

OUTPUT_DIR = PROJECT_ROOT / "docs" / "presentation_pack"
OUTPUT_STEM = f"Technocops_DDC_Product_Overview_v{APP_VERSION.replace('.', '_')}"
DOCX_PATH = OUTPUT_DIR / f"{OUTPUT_STEM}.docx"
PDF_PATH = OUTPUT_DIR / f"{OUTPUT_STEM}.pdf"
LOGO_PATH = PROJECT_ROOT / "assets" / "branding" / "technocops_app_icon.png"

PRIMARY = RGBColor(0x17, 0x39, 0x6D)
ACCENT = RGBColor(0x23, 0x84, 0xFF)
MUTED = RGBColor(0x5A, 0x6B, 0x81)
LIGHT_FILL = "EAF2FF"
PDF_PRIMARY = colors.HexColor("#17396D")
PDF_ACCENT = colors.HexColor("#2384FF")
PDF_LIGHT = colors.HexColor("#EAF2FF")
PDF_MUTED = colors.HexColor("#5A6B81")


@dataclass(frozen=True)
class TableBlock:
    headers: list[str]
    rows: list[list[str]]


@dataclass(frozen=True)
class Section:
    title: str
    items: list[tuple[str, object]]


def main() -> int:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    generated_on = datetime.now().strftime("%d %B %Y")
    sections = build_sections(generated_on)
    write_docx(sections, generated_on)
    write_pdf(sections, generated_on)
    print(DOCX_PATH)
    print(PDF_PATH)
    return 0


def build_sections(generated_on: str) -> list[Section]:
    section_titles = [
        "Executive Summary",
        "Product Purpose and Use Cases",
        "Operating Modes",
        "User Interface Overview",
        "Production Workflow",
        "Core Functionalities",
        "Production-Level Actions",
        "Finalizer-Level Actions",
        "Auto Update and User Experience",
        "Security, Licensing and Data Safety",
        "Technology Stack and Codebase Layout",
        "Examples of Real Transformations",
        "Build, Packaging and Release Process",
        "Recommended Future Enhancements",
        "Presentation Close",
    ]

    return [
        Section(
            "Document Guide",
            [
                (
                    "paragraph",
                    f"This presentation document was generated on {generated_on} for {APP_NAME} {APP_VERSION_LABEL}. It is designed for product, client, and technical presentation use.",
                ),
                ("bullets", section_titles),
            ],
        ),
        Section(
            "Executive Summary",
            [
                (
                    "paragraph",
                    f"{APP_NAME} is a Windows desktop production tool that converts OCR-generated HTML into structured DTBook XML, validates the result, finalizes IDs when needed, and now supports a one-click automatic update flow.",
                ),
                (
                    "bullets",
                    [
                        "Built for publishing and accessible content production teams that need consistent DTBook output from ABBYY FineReader or similar HTML sources.",
                        "Supports HTML files, folders, ZIP archives, drag-and-drop intake, ordered batching, metadata capture, validation, export, and post-processing.",
                        "Provides a dedicated XML Finalizer mode for safe ID regeneration without changing the document content.",
                        "Includes machine-bound licensing, startup security checks, release packaging, and silent in-app updater support.",
                        "Latest production workflow saves XML inside an automatically created output folder and no longer generates report.json.",
                    ],
                ),
            ],
        ),
        Section(
            "Product Purpose and Use Cases",
            [
                (
                    "paragraph",
                    "The tool solves the gap between OCR output and production-ready DTBook XML. It reduces manual cleanup, enforces consistent metadata, normalizes structure, and gives operators a guided review workflow before delivery.",
                ),
                (
                    "bullets",
                    [
                        "Convert multiple chapter HTML files into one DTBook package.",
                        "Normalize OCR HTML into semantic DTBook tags such as level1-level6, pagenum, list, sidebar, linegroup, imggroup, caption, and table.",
                        "Apply metadata required by production workflows, including UID, identifier, title, dates, producer, source publisher, ISBN, and document type.",
                        "Validate output before export and block save when critical issues exist, unless the operator explicitly disables the safety block.",
                        "Repair page IDs and level IDs for existing XML files without reopening the original HTML source files.",
                    ],
                ),
            ],
        ),
        Section(
            "Operating Modes",
            [
                (
                    "table",
                    TableBlock(
                        headers=["Mode", "Input", "Main Goal", "Output"],
                        rows=[
                            [
                                "Production Conversion",
                                "HTML files, folders, or ZIP archives",
                                "Generate new DTBook XML from source content",
                                "XML file, text validation report, copied image folder",
                            ],
                            [
                                "XML Finalizer",
                                "Existing XML file",
                                "Regenerate page IDs and/or level IDs without changing content",
                                "Updated XML preview ready to save",
                            ],
                        ],
                    ),
                ),
                (
                    "paragraph",
                    "Production Conversion is the main authoring and export path. XML Finalizer is a controlled repair path for already-generated XML where only identifiers need to be normalized.",
                ),
            ],
        ),
        Section(
            "User Interface Overview",
            [
                (
                    "table",
                    TableBlock(
                        headers=["UI Area", "Purpose", "Key Controls"],
                        rows=[
                            ["Header", "Branding, status, licensing, update access", "Check Updates, About, status badge, trial/activation status"],
                            ["Input Sources", "Collect and order source files", "Add HTML, Add ZIP, Add Folder, drag-and-drop, Move Up, Move Down, Remove, Clear"],
                            ["DTBook Metadata", "Capture required production metadata", "UID, identifier, title, language, document type, dates, authors, ISBN, producer"],
                            ["Conversion", "Run conversion and review status", "Generate XML, Save XML, page range start, critical error export block"],
                            ["ID Finalizer", "Post-process existing XML safely", "Load XML, Regenerate Page IDs, Regenerate Level IDs, Apply ID Finalizer"],
                            ["Preview Tabs", "Let operator inspect source and result", "Input Preview, XML Preview, Logs"],
                        ],
                    ),
                ),
                (
                    "bullets",
                    [
                        "The interface uses a branded dark desktop layout built with PyQt6 and custom styles.",
                        "Input preview and XML preview help operators quickly compare source and result without leaving the app.",
                        "Progress messages and a progress bar make long conversions easier to monitor during production work.",
                    ],
                ),
            ],
        ),
        Section(
            "Production Workflow",
            [
                (
                    "table",
                    TableBlock(
                        headers=["Step", "Operator Action", "System Action", "Result"],
                        rows=[
                            ["1", "Load HTML, ZIP, or folder", "Collect files, preserve order, extract ZIPs into temporary working space", "Input queue is ready"],
                            ["2", "Review metadata", "Apply defaults, detect language, suggest title/ISBN/publisher", "Metadata form is pre-filled where possible"],
                            ["3", "Set generated page start", "Store page numbering baseline", "All page markers continue from the selected start value"],
                            ["4", "Generate XML", "Parse HTML, convert structure, attach images, validate output", "XML preview and logs are populated"],
                            ["5", "Review issues", "Sort warnings/errors by severity and location", "Operator can decide whether to export"],
                            ["6", "Save XML", "Create output folder automatically, save XML, save text report, copy image assets", "Deliverable package is written beside the source"],
                        ],
                    ),
                ),
                (
                    "paragraph",
                    "The current save behavior is fully automated: when the operator clicks Save XML, the application creates an output folder beside the source file and stores the XML there. The output package now includes the XML, a text validation report, and the img folder when images are used.",
                ),
            ],
        ),
        Section(
            "Core Functionalities",
            [
                (
                    "heading2",
                    "Input and Batch Handling",
                ),
                (
                    "bullets",
                    [
                        "Supports .html and .htm files directly.",
                        "Supports ZIP archives and folder intake.",
                        "Preserves sequence and allows manual reorder before conversion.",
                        "Reads source HTML using fallback encodings such as utf-8, utf-8-sig, cp1252, and latin-1.",
                    ],
                ),
                (
                    "heading2",
                    "Metadata and Document Intelligence",
                ),
                (
                    "bullets",
                    [
                        "Fixed production publisher is enforced as Dedicon.",
                        "Supports document type, source publisher, producer, ISBN, completion date, produced date, and multi-author entry.",
                        "Generate IDs button derives UID and Identifier intelligently when possible.",
                        "Metadata extractor suggests title, ISBN, and publisher from HTML content.",
                        "Language detector checks declared language first, then uses stopword heuristics for English, Dutch, Swedish, and Romanian.",
                    ],
                ),
                (
                    "heading2",
                    "Conversion Engine",
                ),
                (
                    "bullets",
                    [
                        "Builds DTBook root, metadata, frontmatter, bodymatter, and rearmatter.",
                        "Converts headings into nested level1-level6 sections and warns on heading jumps.",
                        "Converts lists, tables, sidebars, poetic line groups, page markers, figures, captions, and inline emphasis.",
                        "Normalizes broken spacing, OCR artifacts, split emphasis, bracketed inline markup, and malformed heading tokens.",
                        "Corrects image-caption order so imggroup always writes the image first and captions after it.",
                    ],
                ),
                (
                    "heading2",
                    "Validation and Export",
                ),
                (
                    "bullets",
                    [
                        "Checks XML well-formedness.",
                        "Confirms required DTBook structural tags and required metadata fields.",
                        "Runs DTD validation using the bundled dtbook-basic.dtd when available.",
                        "Writes a human-readable .report.txt file for export review.",
                        "No longer generates report.json during save.",
                    ],
                ),
            ],
        ),
        Section(
            "Production-Level Actions",
            [
                (
                    "table",
                    TableBlock(
                        headers=["User Action", "Internal Processing", "Production Benefit"],
                        rows=[
                            ["Add ZIP", "Extract archive, discover HTML files, sort naturally, keep archive as save anchor", "Operators can receive a single client package and still save outputs beside the original ZIP"],
                            ["Generate XML", "Run DTBookConverter, then DTBookValidator", "One-button creation plus quality gate"],
                            ["Block export on critical errors", "Disable save when critical validation issues exist", "Prevents accidental delivery of structurally broken XML"],
                            ["Save XML", "Create output folder, name file from dtb:uid or metadata, copy images into img", "Consistent delivery package every time"],
                            ["Check Updates", "Call GitHub latest release API and compare semantic version", "Keeps desktop installations aligned with current packaged build"],
                        ],
                    ),
                ),
            ],
        ),
        Section(
            "Finalizer-Level Actions",
            [
                (
                    "paragraph",
                    "The XML Finalizer is intentionally narrow in scope. It is designed for safe post-processing after conversion or for already existing XML files that only need identifier cleanup.",
                ),
                (
                    "table",
                    TableBlock(
                        headers=["Finalizer Action", "What Changes", "What Does Not Change"],
                        rows=[
                            ["Regenerate Page IDs", "Updates pagenum id and page attributes based on visible page values", "Content text, structure, headings, captions, and order remain unchanged"],
                            ["Regenerate Level IDs", "Renumbers nested level IDs sequentially and normalizes closing tag depth", "Heading text and surrounding content remain unchanged"],
                            ["Load Existing XML", "Loads XML directly into preview without reconversion", "Original HTML is not required"],
                            ["Apply ID Finalizer", "Creates a revised XML preview ready to save", "No semantic content rewrite occurs"],
                        ],
                    ),
                ),
                (
                    "bullets",
                    [
                        "This mode is ideal when production teams receive an XML file with inconsistent IDs but the content itself is already approved.",
                        "Because only identifiers are touched, the risk of accidental content drift is much lower than rerunning a full HTML conversion.",
                    ],
                ),
            ],
        ),
        Section(
            "Auto Update and User Experience",
            [
                (
                    "paragraph",
                    "The application includes an in-app updater that checks the latest GitHub release, downloads the preferred installer asset, launches a silent Windows update, and restarts the application automatically.",
                ),
                (
                    "bullets",
                    [
                        "Update checks are based on the configured GitHub repository and compare the current version with the latest release tag.",
                        "The update prompt presents Update Now and Later instead of sending the user to GitHub.",
                        "When the user starts an update, the app shows Update is in progress messaging and a live download/progress state.",
                        "The background PowerShell helper waits for the app to close, runs the installer silently, and restarts the app when complete.",
                        "Local secure storage is kept under LOCALAPPDATA, so licensing and updater data are preserved during normal upgrades.",
                    ],
                ),
                (
                    "table",
                    TableBlock(
                        headers=["Updater Stage", "System Behavior"],
                        rows=[
                            ["Silent check", "Reads latest release metadata from GitHub Releases API"],
                            ["Decision", "Compares latest tag against APP_VERSION"],
                            ["Download", "Selects preferred installer asset (.exe or .msi) and streams it to secure app data"],
                            ["Install", "Starts a hidden PowerShell updater script with silent installer arguments"],
                            ["Restart", "Relaunches the desktop app after a successful install"],
                        ],
                    ),
                ),
            ],
        ),
        Section(
            "Security, Licensing and Data Safety",
            [
                (
                    "bullets",
                    [
                        "Startup security checks detect attached debuggers and verify asset integrity against a generated SHA-256 manifest.",
                        "License state is stored in secure local storage protected with Windows DPAPI and mirrored to the Windows Registry.",
                        "Trial flow is machine-bound and time-bound, with activation keys derived from the machine ID.",
                        "Update packages download into secure app data rather than mixing with user content folders.",
                        "Normal upgrades do not target the output XML folders, so operator-produced XML deliverables remain separate from application state.",
                    ],
                ),
            ],
        ),
        Section(
            "Technology Stack and Codebase Layout",
            [
                (
                    "table",
                    TableBlock(
                        headers=["Category", "Used In This Tool"],
                        rows=[
                            ["Primary programming language", "Python 3 desktop application"],
                            ["UI framework", "PyQt6"],
                            ["XML/HTML parsing", "lxml plus regex-based cleanup"],
                            ["HTTP and update checks", "requests"],
                            ["Windows integration", "ctypes, winreg, PowerShell"],
                            ["Packaging", "PyInstaller, Inno Setup, Compress-Archive"],
                            ["Data and file handling", "pathlib, shutil, zipfile, tempfile, json, hashlib, uuid"],
                            ["Markup and content types", "HTML input, DTBook XML output, DTD validation"],
                        ],
                    ),
                ),
                (
                    "table",
                    TableBlock(
                        headers=["Module", "Responsibility"],
                        rows=[
                            ["app.py", "Bootstrap, splash screen, security checks, license gate, main window launch"],
                            ["ui/main_window.py", "Main workflow controller and user interactions"],
                            ["ui/widgets.py", "Metadata editor, page range, finalizer controls, summary cards, drag-drop list"],
                            ["services/dtbook_converter.py", "Rule-based HTML to DTBook transformation engine"],
                            ["services/conversion_service.py", "Conversion orchestration, saving, ID finalizer"],
                            ["services/validation.py", "DTBook structural checks and DTD validation"],
                            ["services/update_service.py", "Release checks, installer download, background update script"],
                            ["services/license_service.py", "Trial, activation, local state integrity, machine binding"],
                            ["services/security_service.py and windows_security.py", "Anti-debug and asset/license protection"],
                            ["build_release.ps1 and installer .iss", "Production packaging, smoke tests, installer creation"],
                        ],
                    ),
                ),
            ],
        ),
        Section(
            "Examples of Real Transformations",
            [
                (
                    "heading2",
                    "Example 1: Inline Formatting",
                ),
                (
                    "code",
                    "HTML input:\n<span style=\"font-weight:bold;\">Bold text</span>\n<span style=\"font-style:italic;\">Italic text</span>\n\nDTBook output:\n<strong>Bold text</strong>\n<em>Italic text</em>",
                ),
                (
                    "heading2",
                    "Example 2: Image and Caption Normalization",
                ),
                (
                    "code",
                    "Input sequence:\nCaption paragraph\nImage node\nSecond caption paragraph\n\nDTBook output:\n<imggroup>\n  <img src=\"img/cover.jpg\" alt=\"afbeelding\"/>\n  <caption>\n    <p><em>Caption paragraph</em></p>\n    <p><em>Second caption paragraph</em></p>\n  </caption>\n</imggroup>",
                ),
                (
                    "heading2",
                    "Example 3: Generated Page Start",
                ),
                (
                    "code",
                    "If the operator sets Generated Page Number Start to 1001,\nthen consecutive HTML page markers become:\n<pagenum id=\"page-1001\" page=\"normal\">1001</pagenum>\n<pagenum id=\"page-1002\" page=\"normal\">1002</pagenum>",
                ),
                (
                    "heading2",
                    "Example 4: Finalizer Repair",
                ),
                (
                    "code",
                    "Before finalizer:\n<level3 id=\"broken-two\"><h3>Two</h3><pagenum id=\"broken-b\" page=\"normal\">1a</pagenum></level3>\n\nAfter finalizer:\n<level2 id=\"l-2\"><h3>Two</h3><pagenum id=\"page-1a\" page=\"special\">1a</pagenum></level2>",
                ),
            ],
        ),
        Section(
            "Build, Packaging and Release Process",
            [
                (
                    "bullets",
                    [
                        "build_release.ps1 generates brand assets and a security manifest before packaging.",
                        "Prebuild smoke checks validate UI defaults, conversion behavior, license handling, security checks, and update service configuration.",
                        "PyInstaller produces the portable Windows bundle.",
                        "Postbuild smoke checks validate the packaged distribution.",
                        "Compress-Archive produces the portable ZIP and Inno Setup compiles the installer EXE.",
                        "The release workflow is compatible with GitHub Releases for end-user update delivery.",
                    ],
                ),
            ],
        ),
        Section(
            "Recommended Future Enhancements",
            [
                (
                    "table",
                    TableBlock(
                        headers=["Priority", "Suggested Enhancement", "Why It Matters"],
                        rows=[
                            ["High", "Batch preset profiles by publisher or client", "Reduce repetitive metadata entry and rule differences across customers"],
                            ["High", "Side-by-side source vs XML diff viewer", "Speeds up operator QA and presentation demos"],
                            ["High", "Validation dashboard with filters and severity grouping", "Makes large batch review faster and easier"],
                            ["Medium", "Watch folder or batch queue automation", "Supports higher-volume production runs"],
                            ["Medium", "Custom rule editor for OCR cleanup patterns", "Lets teams tune transformations without editing source code"],
                            ["Medium", "Signed installers and release signature verification", "Improves trust, enterprise deployment, and update safety"],
                            ["Medium", "Stable and beta update channels", "Allows controlled rollouts for new features"],
                            ["Low", "Expanded multilingual metadata templates", "Improves onboarding for additional language markets"],
                            ["Low", "Export-ready QA summary report for clients", "Helps communicate validation status externally"],
                            ["Low", "Audit trail of conversion settings and finalizer actions", "Improves reproducibility and compliance tracking"],
                        ],
                    ),
                ),
            ],
        ),
        Section(
            "Presentation Close",
            [
                (
                    "bullets",
                    [
                        f"{APP_NAME} is already positioned as a production-focused DTBook preparation desktop application, not just a simple converter.",
                        "Its strongest value is the combination of operator-friendly UI, deterministic conversion rules, validation gates, and safe post-processing.",
                        "The new save behavior, image-caption correction, and automatic update flow improve real production reliability.",
                        "The next major growth area is advanced QA workflow, reusable presets, and enterprise-ready deployment controls.",
                    ],
                ),
            ],
        ),
    ]


def write_docx(sections: list[Section], generated_on: str) -> None:
    document = Document()
    configure_docx_document(document)
    add_docx_cover(document, generated_on)

    for section in sections:
        document.add_page_break()
        heading = document.add_paragraph()
        heading.style = document.styles["Heading 1"]
        heading.add_run(section.title)
        for item_type, payload in section.items:
            add_docx_item(document, item_type, payload)

    document.core_properties.author = COMPANY_NAME
    document.core_properties.company = COMPANY_NAME
    document.core_properties.title = f"{APP_NAME} Product Overview"
    document.save(DOCX_PATH)


def configure_docx_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "Calibri"
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    heading1.font.color.rgb = PRIMARY

    heading2 = document.styles["Heading 2"]
    heading2.font.name = "Calibri"
    heading2.font.size = Pt(13)
    heading2.font.bold = True
    heading2.font.color.rgb = ACCENT

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f"{APP_NAME} {APP_VERSION_LABEL} | {COMPANY_NAME}")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = MUTED


def add_docx_cover(document: Document, generated_on: str) -> None:
    cover = document.sections[0]
    cover.start_type = WD_SECTION.NEW_PAGE

    if LOGO_PATH.exists():
        image_paragraph = document.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.add_run().add_picture(str(LOGO_PATH), width=Inches(1.05))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"{APP_NAME}\nProduct Overview and Presentation Pack")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = PRIMARY

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "Functionality, architecture, production workflow, XML finalizer behavior, security model, and auto update flow"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = MUTED

    details = document.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details_run = details.add_run(f"Version {APP_VERSION_LABEL} | Generated on {generated_on}\nPrepared for next-day presentation use")
    details_run.font.size = Pt(11)
    details_run.font.color.rgb = ACCENT


def add_docx_item(document: Document, item_type: str, payload: object) -> None:
    if item_type == "paragraph":
        paragraph = document.add_paragraph(str(payload))
        paragraph.paragraph_format.space_after = Pt(6)
        return

    if item_type == "bullets":
        for bullet in payload:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(str(bullet))
            paragraph.paragraph_format.space_after = Pt(2)
        return

    if item_type == "heading2":
        paragraph = document.add_paragraph()
        paragraph.style = document.styles["Heading 2"]
        paragraph.add_run(str(payload))
        return

    if item_type == "table":
        add_docx_table(document, payload)
        return

    if item_type == "code":
        add_docx_code_block(document, str(payload))
        return

    raise ValueError(f"Unsupported DOCX item type: {item_type}")


def add_docx_table(document: Document, block: TableBlock) -> None:
    table = document.add_table(rows=1, cols=len(block.headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for index, header in enumerate(block.headers):
        cell = header_cells[index]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_docx_cell(cell, "D9E8FF")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = PRIMARY

    for row_values in block.rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row_values):
            row_cells[index].text = value
            row_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document.add_paragraph()


def add_docx_code_block(document: Document, code_text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_docx_cell(cell, LIGHT_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(code_text.splitlines()):
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        if index < len(code_text.splitlines()) - 1:
            run.add_break()
    document.add_paragraph()


def shade_docx_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def write_pdf(sections: list[Section], generated_on: str) -> None:
    styles = build_pdf_styles()
    story = []

    if LOGO_PATH.exists():
        story.append(Image(str(LOGO_PATH), width=0.9 * inch, height=0.9 * inch))
        story.append(Spacer(1, 0.18 * inch))

    story.append(Paragraph(APP_NAME, styles["CoverTitle"]))
    story.append(Paragraph("Product Overview and Presentation Pack", styles["CoverSubTitle"]))
    story.append(
        Paragraph(
            "Functionality, architecture, production workflow, XML finalizer behavior, security model, and auto update flow",
            styles["CoverLead"],
        )
    )
    story.append(Spacer(1, 0.16 * inch))
    story.append(Paragraph(f"Version {APP_VERSION_LABEL} | Generated on {generated_on}", styles["Meta"]))
    story.append(Paragraph("Prepared for next-day presentation use", styles["Meta"]))
    story.append(PageBreak())

    for index, section in enumerate(sections):
        if index:
            story.append(Spacer(1, 0.08 * inch))
        story.append(Paragraph(section.title, styles["Heading1"]))
        for item_type, payload in section.items:
            add_pdf_item(story, styles, item_type, payload)

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.55 * inch,
        title=f"{APP_NAME} Product Overview",
        author=COMPANY_NAME,
    )
    doc.build(story, onFirstPage=pdf_canvas_footer, onLaterPages=pdf_canvas_footer)


def build_pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CoverTitle",
            parent=styles["Title"],
            alignment=TA_CENTER,
            textColor=PDF_PRIMARY,
            fontSize=22,
            leading=26,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverSubTitle",
            parent=styles["Heading1"],
            alignment=TA_CENTER,
            textColor=PDF_ACCENT,
            fontSize=15,
            leading=18,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverLead",
            parent=styles["BodyText"],
            alignment=TA_CENTER,
            textColor=PDF_MUTED,
            fontSize=10.5,
            leading=14,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Meta",
            parent=styles["BodyText"],
            alignment=TA_CENTER,
            textColor=PDF_MUTED,
            fontSize=9.5,
            leading=12,
            spaceAfter=3,
        )
    )
    styles["Heading1"].textColor = PDF_PRIMARY
    styles["Heading1"].fontSize = 16
    styles["Heading1"].leading = 20
    styles["Heading1"].spaceAfter = 8
    styles["Heading2"].textColor = PDF_ACCENT
    styles["Heading2"].fontSize = 12
    styles["Heading2"].leading = 15
    styles["Heading2"].spaceAfter = 6
    styles["BodyText"].fontSize = 9.5
    styles["BodyText"].leading = 13
    styles["BodyText"].spaceAfter = 5
    styles.add(
        ParagraphStyle(
            name="CodeBlock",
            parent=styles["BodyText"],
            fontName="Courier",
            backColor=PDF_LIGHT,
            borderPadding=6,
            borderColor=colors.HexColor("#B9D2F5"),
            borderWidth=0.5,
            leading=11,
            spaceBefore=3,
            spaceAfter=8,
        )
    )
    return styles


def add_pdf_item(story, styles, item_type: str, payload: object) -> None:
    if item_type == "paragraph":
        story.append(Paragraph(str(payload), styles["BodyText"]))
        return

    if item_type == "bullets":
        items = [
            ListItem(Paragraph(str(value), styles["BodyText"]), leftIndent=12)
            for value in payload
        ]
        story.append(ListFlowable(items, bulletType="bullet", start="circle", leftIndent=16))
        story.append(Spacer(1, 0.08 * inch))
        return

    if item_type == "heading2":
        story.append(Paragraph(str(payload), styles["Heading2"]))
        return

    if item_type == "table":
        add_pdf_table(story, styles, payload)
        return

    if item_type == "code":
        story.append(Preformatted(str(payload), styles["CodeBlock"]))
        return

    raise ValueError(f"Unsupported PDF item type: {item_type}")


def add_pdf_table(story, styles, block: TableBlock) -> None:
    data = [block.headers] + block.rows
    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), PDF_LIGHT),
                ("TEXTCOLOR", (0, 0), (-1, 0), PDF_PRIMARY),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("LEADING", (0, 0), (-1, -1), 10),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#BDD2F0")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FBFF")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 0.12 * inch))


def pdf_canvas_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(PDF_MUTED)
    canvas.drawCentredString(A4[0] / 2, 20, f"{APP_NAME} {APP_VERSION_LABEL} | {COMPANY_NAME}")
    canvas.restoreState()


if __name__ == "__main__":
    raise SystemExit(main())
